import os
import io
import tempfile
import random
from datetime import date
from functools import wraps
from flask import Flask, jsonify, request, render_template, send_file, redirect, url_for, session
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.shared import Cm
from docx2pdf import convert
from flask_sqlalchemy import SQLAlchemy
from sqlalchemy.exc import OperationalError, IntegrityError

# Import db and all models from the models.py file
from models import db, Employee, User, Leave 

app = Flask(__name__)
app.secret_key = 'your_super_secret_key'

# Configure the app
app.config["SQLALCHEMY_DATABASE_URI"] = "sqlite:///hire.db"
app.config["SQLALCHEMY_TRACK_MODIFICATIONS"] = False

# Initialize the db object with the app
with app.app_context():
    db.init_app(app)

@app.template_filter('zfill')
def zfill_filter(s, width):
    return str(s).zfill(width)

POSITIONS = ["Manager", "Developer", "Data Analyst", "HR", "Designer", "Intern"]

def serialize_emp(e):
    return {
        "id": e.id,
        "name": e.name,
        "age": e.age,
        "salary": e.salary,
        "years_in_job": e.years_in_job,
        "position": e.position,
    }

def seed_data(total=10000, batch_size=1000):
    print(f"Seeding {total} employees...")
    created = 0
    while created < total:
        batch = []
        start_id = created + 1
        end_id = min(created + batch_size, total)
        for i in range(start_id, end_id + 1):
            batch.append(
                Employee(
                    id=i,
                    name=f"Employee_{i}",
                    age=random.randint(20, 60),
                    salary=random.randint(20000, 120000),
                    years_in_job=random.randint(0, 20),
                    position=random.choice(POSITIONS),
                )
            )
        db.session.bulk_save_objects(batch)
        db.session.commit()
        created = end_id
        print(f"  -> {created}/{total} inserted")

    print("Seeding default user...")
    admin_user = User(username='admin', password='password123')
    db.session.add(admin_user)
    db.session.commit()
    print("Default user 'admin' created.")

def login_required(f):
    @wraps(f)
    def wrap(*args, **kwargs):
        if 'logged_in' in session:
            return f(*args, **kwargs)
        else:
            return redirect(url_for('login'))
    return wrap

@app.route("/login", methods=["GET", "POST"])
def login():
    if request.method == "POST":
        username = request.form.get("username")
        password = request.form.get("password")
        user = User.query.filter_by(username=username, password=password).first()
        if user:
            session['logged_in'] = True
            session['username'] = user.username
            return redirect(url_for('list_employees'))
        else:
            return render_template("login.html", error="Invalid username or password")
    return render_template("login.html")

@app.route("/logout")
def logout():
    session.clear()
    return redirect(url_for('login'))

@app.route("/", methods=["GET"])
def root():
    return redirect(url_for('login'))

@app.route("/employees", methods=["GET"])
@login_required
def list_employees():
    q = db.session.query(Employee)
    position = request.args.get("position")
    min_salary = request.args.get("min_salary", type=int)
    max_salary = request.args.get("max_salary", type=int)

    if position:
        q = q.filter(Employee.position == position)
    if min_salary is not None:
        q = q.filter(Employee.salary >= min_salary)
    if max_salary is not None:
        q = q.filter(Employee.salary <= max_salary)

    page = request.args.get("page", default=1, type=int)
    per_page = request.args.get("per_page", default=50, type=int)
    per_page = max(1, min(per_page, 200))

    total = q.count()
    rows = q.order_by(Employee.id).offset((page - 1) * per_page).limit(per_page).all()

    return render_template("employees.html",
                           total=total,
                           page=page,
                           per_page=per_page,
                           data=[serialize_emp(e) for e in rows])

@app.route("/employee/<int:emp_id>", methods=["GET"])
@login_required
def get_employee(emp_id):
    e = db.session.get(Employee, emp_id)
    if not e:
        return jsonify({"error": "Employee not found"}), 404
    return jsonify(serialize_emp(e))

@app.route("/employees", methods=["POST"])
@login_required
def add_employee():
    data = request.json
    try:
        new_employee = Employee(
            name=data['name'],
            age=data['age'],
            salary=data['salary'],
            years_in_job=data['years_in_job'],
            position=data['position']
        )
        db.session.add(new_employee)
        db.session.commit()
        return jsonify({"message": "Employee added successfully!", "id": new_employee.id}), 201
    except KeyError as e:
        return jsonify({"error": f"Missing field: {e}"}), 400
    except Exception as e:
        db.session.rollback()
        return jsonify({"error": str(e)}), 500

@app.route("/positions", methods=["GET"])
@login_required
def get_positions():
    return jsonify(sorted(POSITIONS))

@app.route("/employee/<int:emp_id>/leave-form", methods=["GET"])
@login_required
def leave_form(emp_id):
    employee = db.session.get(Employee, emp_id)
    if not employee:
        return jsonify({"error": "Employee not found"}), 404

    today_date = date.today().strftime("%Y-%m-%d")

    leaves = Leave.query.filter_by(employee_id=emp_id).all()
    leave_data = {
        'total_casual': 0, 'taken_casual': 0,
        'total_sick': 0, 'taken_sick': 0,
        'total_nopay': 0, 'taken_nopay': 0
    }
    for leave in leaves:
        if leave.leave_type == 'Casual Leave':
            leave_data['total_casual'] = 10
            leave_data['taken_casual'] += 1
        elif leave.leave_type == 'Sick Leave':
            leave_data['total_sick'] = 10
            leave_data['taken_sick'] += 1
        elif leave.leave_type == 'Without Pay':
            leave_data['total_nopay'] = 10
            leave_data['taken_nopay'] += 1

    return render_template("leave_application.html", employee=employee, today_date=today_date, leave_data=leave_data)

@app.route("/generate-leave-doc", methods=["POST"])
@login_required
def generate_leave_doc():
    employee_id = request.form.get("employee_id", type=int)
    employee = db.session.get(Employee, employee_id)
    if not employee:
        return jsonify({"error": "Employee not found"}), 404

    start_date = request.form.get("start_date")
    end_date = request.form.get("end_date")
    leave_type = request.form.get("leave_type")
    explanation = request.form.get("explanation")
    mobile = request.form.get("mobile")
    emergency_contact = request.form.get("emergency_contact")
    last_leave_date = request.form.get("last_leave_date")

    total_casual = request.form.get("total_casual")
    taken_casual = request.form.get("taken_casual")
    total_sick = request.form.get("total_sick")
    taken_sick = request.form.get("taken_sick")
    total_nopay = request.form.get("total_nopay")
    taken_nopay = request.form.get("taken_nopay")

    new_leave = Leave(
        employee_id=employee.id,
        leave_type=leave_type,
        start_date=start_date,
        end_date=end_date,
        explanation=explanation
    )
    db.session.add(new_leave)
    db.session.commit()

    doc = Document()
    doc.add_heading("DEPARTMENT OF STATISTICS,ISLAIMC UNIVERSITY", level=1)
    doc.add_paragraph("APPLICATION FOR LEAVE", style='Heading 2')

    doc.add_paragraph("\n")
    doc.add_paragraph(f"Date: {date.today().strftime('%d-%b-%Y')}")

    doc.add_paragraph("\n")
    doc.add_paragraph("Personal Info", style='Heading 3')
    doc.add_paragraph(f"Name: {employee.name}")
    doc.add_paragraph(f"Employee ID: T{str(employee.id).zfill(6)}")
    doc.add_paragraph(f"Designation: {employee.position}")
    doc.add_paragraph(f"Mobile: {mobile}")

    doc.add_paragraph("\n")
    doc.add_paragraph("Leave Request", style='Heading 3')
    doc.add_paragraph(f"Date: {start_date} To {end_date}")
    doc.add_paragraph(f"Leave Type: {leave_type}")
    doc.add_paragraph(f"Last Leave Date: {last_leave_date}")
    doc.add_paragraph(f"Explanation: {explanation}")
    doc.add_paragraph(f"Emergency Contact: {emergency_contact}")

    doc.add_paragraph("\n")
    doc.add_paragraph("Leave Details", style='Heading 3')
    table = doc.add_table(rows=1, cols=3, style='Table Grid')
    table.cell(0, 0).text = "Leave Type"
    table.cell(0, 1).text = "Total"
    table.cell(0, 2).text = "Taken"

    row_cells = table.add_row().cells
    row_cells[0].text = "Casual Leave"
    row_cells[1].text = total_casual
    row_cells[2].text = taken_casual

    row_cells = table.add_row().cells
    row_cells[0].text = "Sick Leave"
    row_cells[1].text = total_sick
    row_cells[2].text = taken_sick

    row_cells = table.add_row().cells
    row_cells[0].text = "Without Pay"
    row_cells[1].text = total_nopay
    row_cells[2].text = taken_nopay

    doc.add_paragraph("\n\n")
    doc.add_paragraph("Applicant's Signature").alignment = WD_ALIGN_PARAGRAPH.RIGHT
    doc.add_paragraph(f"Date: {date.today().strftime('%d/%m/%Y')}").alignment = WD_ALIGN_PARAGRAPH.RIGHT

    with tempfile.NamedTemporaryFile(delete=False, suffix=".docx") as temp_docx_file:
        doc.save(temp_docx_file.name)
        temp_docx_path = temp_docx_file.name

    temp_pdf_path = os.path.join(tempfile.gettempdir(), f'Leave_Application_{employee.name}.pdf')
    os.makedirs(os.path.dirname(temp_pdf_path), exist_ok=True)

    try:
        convert(temp_docx_path, temp_pdf_path)
    except Exception as e:
        print(f"Conversion failed: {e}")
        return jsonify({"error": "Failed to convert document to PDF."}), 500

    response = send_file(
        temp_pdf_path,
        mimetype='application/pdf',
        as_attachment=True,
        download_name=f'Leave_Application_{employee.name}.pdf'
    )

    os.remove(temp_docx_path)

    return response

if __name__ == "__main__":
    import argparse
    parser = argparse.ArgumentParser()
    parser.add_argument("--init-db", action="store_true", help="Create tables and seed 10k employees (runs once).")
    args = parser.parse_args()

    if args.init_db:
        with app.app_context():
            try:
                db.create_all()
                print("Tables created successfully using create_all().")
            except Exception as e:
                print(f"Failed to create tables with create_all(): {e}")
                print("Attempting to create tables manually...")
                
                db.drop_all() 
                
                Employee.__table__.create(db.engine)
                User.__table__.create(db.engine)
                Leave.__table__.create(db.engine)
                print("Tables created manually.")

            if db.session.query(Employee).count() == 0:
                seed_data(total=10000, batch_size=1000)
                print("Seeding done.")
            else:
                print("Employees already exist. Skipping seeding.")
                
            if not User.query.filter_by(username='admin').first():
                db.session.add(User(username='admin', password='password123'))
                db.session.commit()
                print("Default user 'admin' created.")

    app.run(debug=True)